from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional, Union

import os
import shutil
import subprocess
import sys

import pdfplumber
import pypdfium2
import pytesseract


@dataclass
class ExtractedDoc:
    filename: str
    text: str
    total_pages: int
    ocr_used: bool = False
    failed_pages: list[int] = field(default_factory=list)


def _extract_plumber_page(page) -> str:
    parts = []
    text = page.extract_text(x_tolerance=1, y_tolerance=1) or ""
    if text.strip():
        parts.append(text.strip())
    words = page.extract_words(x_tolerance=1, y_tolerance=1) or []
    if words:
        word_lines = []
        current_line = []
        current_top = None
        for w in words:
            top = w.get("top")
            if current_top is None or (top is not None and abs(top - current_top) > 2):
                if current_line:
                    word_lines.append(" ".join(current_line))
                current_line = [w["text"]]
                current_top = top
            else:
                current_line.append(w["text"])
        if current_line:
            word_lines.append(" ".join(current_line))
        parts.append("\n".join(word_lines))
    tables = page.extract_tables() or []
    if tables:
        for table in tables:
            rows = [
                " | ".join(str(c) if c else "" for c in row)
                for row in table
                if row
            ]
            parts.append("\n".join(rows))
    return "\n\n".join(parts)


def _bundle_base() -> Optional[str]:
    """Return the PyInstaller bundle base dir when frozen, else None."""
    if getattr(sys, "frozen", False):
        meipass = getattr(sys, "_MEIPASS", None)
        if meipass:
            return meipass
        return os.path.dirname(sys.executable)
    return None


def _search_tree(root: str, name: str) -> Optional[str]:
    """Walk a directory tree for an executable file named `name`."""
    try:
        for dirpath, _dirnames, filenames in os.walk(root):
            if name in filenames:
                p = Path(dirpath) / name
                if p.is_file() and (p.stat().st_mode & 0o111):
                    return str(p)
    except Exception:  # noqa: BLE001
        pass
    return None


def find_tesseract() -> Optional[str]:
    candidates = []
    base = _bundle_base()
    if base:
        candidates.append(os.path.join(base, "tesseract", "tesseract"))
        candidates.append(os.path.join(base, "tesseract", "tesseract.exe"))
        candidates.append(os.path.join(base, "tesseract"))
        # PyInstaller may relocate Mach-O binaries into Contents/Frameworks
        candidates.append(os.path.join(base, "..", "Frameworks", "tesseract", "tesseract"))
        candidates.append(_search_tree(base, "tesseract"))
        if getattr(sys, "frozen", False):
            candidates.append(_search_tree(os.path.dirname(sys.executable), "tesseract"))
    if sys.platform == "win32":
        env_path = os.environ.get("TESSERACT_PATH")
        if env_path:
            candidates.append(env_path)
        candidates.extend([
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        ])
    else:
        candidates.extend([
            "/usr/local/bin/tesseract",
            "/opt/homebrew/bin/tesseract",
        ])
    for c in candidates:
        p = Path(c)
        if p.is_file() and (p.stat().st_mode & 0o111):
            return str(p)
    try:
        result = shutil.which("tesseract")
        if result:
            return result
    except Exception:
        pass
    return None


def find_tessdata() -> Optional[str]:
    """Locate the tessdata directory (for TESSDATA_PREFIX)."""
    base = _bundle_base()
    if base:
        cand = os.path.join(base, "tessdata")
        if os.path.isdir(cand):
            return cand
        found = _search_tree(base, "eng.traineddata")
        if found:
            return str(Path(found).parent)
    candidates = [
        "/opt/homebrew/share/tessdata",
        "/usr/local/share/tessdata",
        "/usr/share/tessdata",
    ]
    for d in candidates:
        if os.path.isdir(d):
            return d
    # Homebrew (Apple Silicon / Intel) may keep tessdata in the Cellar
    try:
        out = subprocess.run(["brew", "--prefix", "tesseract"],
                             capture_output=True, text=True).stdout.strip()
        if out and os.path.isdir(os.path.join(out, "share", "tessdata")):
            return os.path.join(out, "share", "tessdata")
    except Exception:  # noqa: BLE001
        pass
    return None


def _ocr_page(pdf_path: Path, page_num: int) -> str:
    pdf = pypdfium2.PdfDocument(str(pdf_path))
    try:
        page = pdf[page_num - 1]
        pil = page.render(scale=400 / 72).to_pil()
    finally:
        pdf.close()
    return pytesseract.image_to_string(pil, lang="eng", config="--psm 6")


def extract_pdf(
    pdf_path: Union[str, Path],
    progress_callback=None,
) -> ExtractedDoc:
    pdf_path = Path(pdf_path)
    tesseract_path = find_tesseract()
    ocr_available = tesseract_path is not None
    if not ocr_available:
        print("[WARN] Tesseract not found – OCR will be skipped; only pdfplumber will be used.")

    if ocr_available:
        pytesseract.pytesseract.tesseract_cmd = tesseract_path
        tessdata = find_tessdata()
        if tessdata:
            os.environ["TESSDATA_PREFIX"] = tessdata

    filename = pdf_path.name
    
    with pdfplumber.open(pdf_path) as pdf:
        total_pages = len(pdf.pages)

    all_text_parts: list[str] = [""] * total_pages
    ocr_used_flag = False
    failed: list[int] = []

    from concurrent.futures import ThreadPoolExecutor

    def process_page(idx: int):
        if progress_callback:
            progress_callback(filename, idx, total_pages)
            
        plumber_text = ""
        try:
            with pdfplumber.open(pdf_path) as pdf:
                page = pdf.pages[idx - 1]
                plumber_text = _extract_plumber_page(page)
        except Exception:
            plumber_text = ""

        if plumber_text.strip() and len(plumber_text.strip()) > 50:
            return idx, f"--- PAGE {idx} ---\n{plumber_text}", False, False

        ocr_text = ""
        ocr_used = False
        if ocr_available:
            try:
                ocr_text = _ocr_page(pdf_path, idx)
                ocr_used = True
            except Exception:
                ocr_text = ""

        combined = plumber_text + ocr_text
        is_failed = len(combined.strip()) < 20
        return idx, f"--- PAGE {idx} ---\n{combined}", ocr_used, is_failed

    max_workers = min(os.cpu_count() or 4, 8)
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = [executor.submit(process_page, idx) for idx in range(1, total_pages + 1)]
        for fut in futures:
            try:
                idx, page_text, ocr_used, is_failed = fut.result()
                all_text_parts[idx - 1] = page_text
                if ocr_used:
                    ocr_used_flag = True
                if is_failed:
                    failed.append(idx)
            except Exception as e:
                print(f"[ERROR] Failed to process page: {e}")

    return ExtractedDoc(
        filename=filename,
        text="\n\n".join(all_text_parts).strip(),
        total_pages=total_pages,
        ocr_used=ocr_used_flag,
        failed_pages=sorted(failed),
    )


# ── Supported file types ──────────────────────────────────────────

SUPPORTED_EXTENSIONS = {
    ".pdf", ".docx", ".doc",
    ".xlsx", ".xls",
    ".csv",
    ".txt", ".md", ".rst", ".rtf",
    ".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".webp",
}


# ── .docx extractor ─────────────────────────────────────────────

def extract_docx(path: Path) -> ExtractedDoc:
    try:
        from docx import Document
    except ImportError:
        return ExtractedDoc(
            filename=path.name, text="", total_pages=0, ocr_used=False,
            failed_pages=[],
        )
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs]
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        tables.append("\n".join(rows))
    text = "\n\n".join(paragraphs)
    if tables:
        text += "\n\n--- TABLES ---\n\n" + "\n\n".join(tables)
    return ExtractedDoc(
        filename=path.name, text=text.strip(), total_pages=1,
        ocr_used=False, failed_pages=[],
    )


# ── .xlsx / .xls extractor ──────────────────────────────────────

def extract_xlsx(path: Path) -> ExtractedDoc:
    try:
        import openpyxl
    except ImportError:
        return ExtractedDoc(
            filename=path.name, text="", total_pages=0, ocr_used=False,
            failed_pages=[],
        )
    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            rows.append(" | ".join(cells))
        text = "\n".join(rows)
        parts.append(f"--- Sheet: {sheet_name} ---\n{text}")
    wb.close()
    return ExtractedDoc(
        filename=path.name, text="\n\n".join(parts).strip(),
        total_pages=1, ocr_used=False, failed_pages=[],
    )


# ── .csv extractor ──────────────────────────────────────────────

def extract_csv(path: Path) -> ExtractedDoc:
    import csv
    rows: list[str] = []
    with open(path, newline="", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        for row in reader:
            rows.append(" | ".join(cell.strip() for cell in row))
    return ExtractedDoc(
        filename=path.name, text="\n".join(rows).strip(),
        total_pages=1, ocr_used=False, failed_pages=[],
    )


# ── .txt / .md extractor ────────────────────────────────────────

def extract_text(path: Path) -> ExtractedDoc:
    with open(path, encoding="utf-8", errors="replace") as f:
        text = f.read()
    return ExtractedDoc(
        filename=path.name, text=text.strip(),
        total_pages=1, ocr_used=False, failed_pages=[],
    )


# ── Image OCR extractor ─────────────────────────────────────────

def extract_image(path: Path) -> ExtractedDoc:
    ocr_path = find_tesseract()
    if not ocr_path:
        return ExtractedDoc(
            filename=path.name, text="", total_pages=0, ocr_used=False,
            failed_pages=[],
        )
    try:
        from PIL import Image
        pytesseract.pytesseract.tesseract_cmd = ocr_path
        img = Image.open(path)
        text = pytesseract.image_to_string(img, lang="eng", config="--psm 6")
        return ExtractedDoc(
            filename=path.name, text=text.strip(), total_pages=1,
            ocr_used=True, failed_pages=[],
        )
    except Exception:
        return ExtractedDoc(
            filename=path.name, text="", total_pages=0, ocr_used=False,
            failed_pages=[],
        )


# ── Dispatcher ──────────────────────────────────────────────────

def extract_file(
    path: Union[str, Path],
    progress_callback=None,
) -> ExtractedDoc:
    path = Path(path)
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_pdf(path, progress_callback)
    elif ext in (".docx", ".doc"):
        return extract_docx(path)
    elif ext in (".xlsx", ".xls"):
        return extract_xlsx(path)
    elif ext == ".csv":
        return extract_csv(path)
    elif ext in (".txt", ".md", ".rst", ".rtf"):
        return extract_text(path)
    elif ext in (".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".webp"):
        return extract_image(path)
    else:
        # Fall back to text extraction if the file has no extension (e.g. "Part 2 Bidder PRoposal ")
        if not ext:
            return extract_text(path)
        raise ValueError(f"Unsupported file type: {ext}")


# ── Directory extractor (all supported types) ───────────────────

def extract_directory(
    dir_path: Union[str, Path],
    progress_callback=None,
) -> list[ExtractedDoc]:
    dir_path = Path(dir_path)
    files: list[Path] = []
    # Scan all files in directory
    for p in dir_path.rglob("*"):
        if p.is_file():
            if p.name.startswith("."):
                continue
            # Accept if it matches a supported extension OR has no extension
            if p.suffix.lower() in SUPPORTED_EXTENSIONS or not p.suffix:
                files.append(p)
    files = sorted(set(files))
    total = len(files)
    results: list[ExtractedDoc] = []
    for i, p in enumerate(files):
        if progress_callback:
            progress_callback(p.name, i + 1, total)
        results.append(extract_file(p, progress_callback))
    return results
